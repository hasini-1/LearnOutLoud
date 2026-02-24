import os
from docx import Document
import PyPDF2
import uvicorn
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from groq import Groq
import re

app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

# Initialize Groq with temperature 0.1 for factual answers
client = Groq(api_key="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")

# -----------------------------
# GLOBAL STATES
# -----------------------------
user_orientation = 0
directions = ["Front", "Right", "Back", "Left"]

current_document_text = ""
current_position = 0
loaded_filename = ""
current_document_name = ""
current_document_type = ""

# Automatically detect user's Documents folder
DOCUMENT_FOLDER = os.path.join(os.path.expanduser("~"), "Documents")

class ChatRequest(BaseModel):
    text: str

# -----------------------------
# LANGUAGE DETECTION
# -----------------------------
def detect_language(text):
    """Native Multilingual Support - Language Detection"""
    # Hindi/Devanagari range
    if re.search(r'[\u0900-\u097F]', text):
        return "hi-IN", "Hindi"
    # Telugu range
    elif re.search(r'[\u0C00-\u0C7F]', text):
        return "te-IN", "Telugu"
    # Kannada range
    elif re.search(r'[\u0C80-\u0CFF]', text):
        return "kn-IN", "Kannada"
    else:
        return "en-US", "English"

# -----------------------------
# SPATIAL LOGIC
# -----------------------------
def update_spatial_logic(command: str):
    global user_orientation

    if "turn right" in command:
        user_orientation = (user_orientation + 1) % 4
        return f"Rotated 90 degrees right. Now facing {directions[user_orientation]}."

    if "turn left" in command:
        user_orientation = (user_orientation - 1) % 4
        return f"Rotated 90 degrees left. Now facing {directions[user_orientation]}."

    if "turn around" in command or "180" in command:
        user_orientation = (user_orientation + 2) % 4
        return "Turned around 180 degrees."

    return None

# -----------------------------
# DOCUMENT EXTRACTION FUNCTIONS
# -----------------------------
def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"[Page {page_num + 1}] " + page_text + "\n\n"
    except Exception as e:
        text = f"Error reading PDF: {str(e)}"
    return text.strip()

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        for para_num, paragraph in enumerate(doc.paragraphs):
            if paragraph.text:
                text += paragraph.text + "\n"
        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
    except Exception as e:
        text = f"Error reading DOCX: {str(e)}"
    return text.strip()

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except UnicodeDecodeError:
        # Try different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
        except:
            text = f"Error reading TXT: Unsupported encoding"
    except Exception as e:
        text = f"Error reading TXT: {str(e)}"
    return text.strip()

def read_document(file_path):
    """Read document based on file extension"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    else:
        return "Unsupported file format"

# -----------------------------
# FIND MATCHING FILE (All formats)
# -----------------------------
def find_matching_document(user_input):
    if not os.path.exists(DOCUMENT_FOLDER):
        return None

    # Get all supported documents
    files = [f for f in os.listdir(DOCUMENT_FOLDER) 
             if f.endswith(('.pdf', '.docx', '.txt'))]
    
    user_input_lower = user_input.lower()

    for file in files:
        name_without_ext = os.path.splitext(file)[0].lower()
        if name_without_ext in user_input_lower or file.lower() in user_input_lower:
            return file

    return None

# -----------------------------
# FORMAT FILE LIST FOR SPEECH
# -----------------------------
def format_file_list(files):
    """Format file list for better speech output"""
    if not files:
        return ""
    
    pdf_files = [f for f in files if f.endswith('.pdf')]
    docx_files = [f for f in files if f.endswith('.docx')]
    txt_files = [f for f in files if f.endswith('.txt')]
    
    file_groups = []
    if pdf_files:
        pdf_names = [os.path.splitext(f)[0] for f in pdf_files]
        file_groups.append(f"{len(pdf_files)} PDF file{'s' if len(pdf_files) > 1 else ''}: {', '.join(pdf_names)}")
    if docx_files:
        docx_names = [os.path.splitext(f)[0] for f in docx_files]
        file_groups.append(f"{len(docx_files)} Word file{'s' if len(docx_files) > 1 else ''}: {', '.join(docx_names)}")
    if txt_files:
        txt_names = [os.path.splitext(f)[0] for f in txt_files]
        file_groups.append(f"{len(txt_files)} Text file{'s' if len(txt_files) > 1 else ''}: {', '.join(txt_names)}")
    
    return " | ".join(file_groups)

# -----------------------------
# MAIN ENDPOINT
# -----------------------------
@app.post("/talk")
async def talk(request: ChatRequest):
    global user_orientation, current_document_text, current_position, loaded_filename, current_document_name, current_document_type

    user_input = request.text.strip().lower()
    print(f"🗣️ User said: {request.text}")
    
    # Detect language for response
    lang_code, lang_name = detect_language(request.text)
    print(f"🌐 Detected language: {lang_name}")

    # 1️⃣ Spatial commands
    spatial_reply = update_spatial_logic(user_input)
    if spatial_reply:
        return {"reply": spatial_reply}

    # 2️⃣ LIST ALL DOCUMENTS
    if any(word in user_input for word in ["list", "files", "documents", "what files", "show files", "upload document"]):
        if not os.path.exists(DOCUMENT_FOLDER):
            return {"reply": "Documents folder not found."}

        files = [f for f in os.listdir(DOCUMENT_FOLDER) 
                if f.endswith(('.pdf', '.docx', '.txt'))]

        if not files:
            return {"reply": "No PDF, Word, or Text documents found in your Documents folder."}

        file_list = format_file_list(files)
        return {"reply": f"I found {len(files)} documents in your folder. {file_list}. Say the filename to load it, then say 'read' to begin."}

    # 3️⃣ LOAD DOCUMENT (by filename)
    matched_file = find_matching_document(user_input)
    
    if matched_file and ("load" in user_input or "open" in user_input or "read" in user_input or matched_file.lower() in user_input):
        path = os.path.join(DOCUMENT_FOLDER, matched_file)
        
        # Extract text based on file type
        full_text = read_document(path)
        
        if full_text.startswith("Error"):
            return {"reply": full_text}
        
        current_document_text = full_text
        current_position = 0
        loaded_filename = matched_file
        current_document_name = os.path.splitext(matched_file)[0]
        current_document_type = os.path.splitext(matched_file)[1][1:].upper()
        
        # Get file size info
        word_count = len(full_text.split())
        char_count = len(full_text)
        
        return {"reply": f"{current_document_name} ({current_document_type}) loaded successfully. It has approximately {word_count} words. Say 'read' to begin reading, 'summary' for summary, or 'pause' to pause."}

    # 4️⃣ READ DOCUMENT (continue reading)
    if "read" in user_input:
        if not current_document_text:
            return {"reply": "No document loaded. First say the filename to load it, for example 'load thesis.pdf' or say 'list documents' to see available files."}

        chunk_size = 600  # Optimal for speech
        chunk = current_document_text[current_position:current_position + chunk_size]
        
        if not chunk:
            return {"reply": f"You have reached the end of {current_document_name}. Say 'read again' to restart, or load another document."}
        
        current_position += chunk_size
        
        # Add progress indicator
        progress = int((current_position / len(current_document_text)) * 100)
        
        return {"reply": f"{chunk} [Progress: {progress}%. Say 'continue' to read more, 'pause' to pause, or 'stop' to end.]"}

    # 5️⃣ CONTINUE READING (alternative command)
    if "continue" in user_input or "more" in user_input:
        if not current_document_text:
            return {"reply": "No document loaded."}
            
        chunk_size = 600
        chunk = current_document_text[current_position:current_position + chunk_size]
        
        if not chunk:
            return {"reply": f"You have reached the end of {current_document_name}."}
        
        current_position += chunk_size
        progress = int((current_position / len(current_document_text)) * 100)
        
        return {"reply": f"{chunk} [Progress: {progress}%]"}

    # 6️⃣ READ AGAIN (restart document)
    if "read again" in user_input or "restart" in user_input or "start over" in user_input:
        if current_document_text:
            current_position = 0
            chunk = current_document_text[:600]
            return {"reply": f"Restarting {current_document_name}. {chunk}"}
        else:
            return {"reply": "No document is currently loaded."}

    # 7️⃣ PAUSE
    if "pause" in user_input:
        return {"reply": "Reading paused. Say 'resume' or 'continue' to continue, or 'stop' to end."}

    # 8️⃣ RESUME
    if "resume" in user_input:
        if not current_document_text:
            return {"reply": "No document loaded."}
        
        if current_position >= len(current_document_text):
            return {"reply": f"You have reached the end of {current_document_name}. Say 'read again' to restart."}
        
        chunk_size = 600
        chunk = current_document_text[current_position:current_position + chunk_size]
        current_position += chunk_size
        progress = int((current_position / len(current_document_text)) * 100)
        
        return {"reply": f"Resuming {current_document_name}. {chunk} [Progress: {progress}%]"}

    # 9️⃣ STOP
    if "stop" in user_input or "end" in user_input:
        # Clear current document
        current_document_text = ""
        current_position = 0
        loaded_filename = ""
        current_document_name = ""
        return {"reply": "Stopped reading. Document unloaded. Say 'list documents' to see available files."}

    # 🔟 SUMMARY
    if "summary" in user_input or "summarize" in user_input or "conclusion" in user_input:
        # Check if we have a document to summarize
        doc_to_summarize = current_document_text
        doc_name = current_document_name if current_document_name else "document"
        
        if not doc_to_summarize:
            # Try to find a document from the command
            matched_file = find_matching_document(user_input)
            if matched_file:
                path = os.path.join(DOCUMENT_FOLDER, matched_file)
                doc_to_summarize = read_document(path)
                doc_name = os.path.splitext(matched_file)[0]
            else:
                return {"reply": "No document loaded or specified. Please load a document first or include the filename, e.g., 'summarize thesis.pdf'."}
        
        if not doc_to_summarize or doc_to_summarize.startswith("Error"):
            return {"reply": "Sorry, I couldn't read the document."}
        
        # Truncate if too long
        if len(doc_to_summarize) > 4000:
            doc_to_summarize = doc_to_summarize[:4000] + "..."
        
        try:
            completion = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": f"Provide a concise summary in {lang_name}. Include main points and key findings."},
                    {"role": "user", "content": f"Please summarize this document: {doc_to_summarize}"}
                ],
                temperature=0.1
            )
            
            summary = completion.choices[0].message.content
            return {"reply": f"Summary of {doc_name}: {summary}"}
            
        except Exception as e:
            return {"reply": f"Sorry, I couldn't generate a summary. Error: {str(e)}"}

    # 1️⃣1️⃣ FIND SPECIFIC INFORMATION
    if "find" in user_input or "search" in user_input or "look for" in user_input:
        if not current_document_text:
            return {"reply": "No document loaded. Please load a document first."}
        
        # Extract search term
        search_term = user_input.replace("find", "").replace("search", "").replace("look for", "").strip()
        
        if not search_term:
            return {"reply": "What would you like me to find?"}
        
        # Search in document
        sentences = current_document_text.split('.')
        matches = [s for s in sentences if search_term.lower() in s.lower()]
        
        if matches:
            return {"reply": f"I found {len(matches)} matches. Here's the first one: {matches[0].strip()}."}
        else:
            return {"reply": f"I couldn't find '{search_term}' in the document."}

    # 1️⃣2️⃣ HELP / INTRODUCTION
    if any(word in user_input for word in ["help", "introduce", "what can you do", "commands"]):
        intro = ("I am your document assistant. I can help you with:\n"
                "📋 'list documents' - Show all PDF, Word, and Text files in your Documents folder\n"
                "📖 'load filename' - Load a specific document (e.g., 'load thesis.pdf')\n"
                "📖 'read' - Start reading the loaded document\n"
                "📖 'continue' - Read more of the current document\n"
                "📝 'summary' - Get an AI summary of the document\n"
                "🔍 'find keyword' - Search for specific words in the document\n"
                "⏸️ 'pause' - Pause reading\n"
                "▶️ 'resume' - Resume reading\n"
                "⏹️ 'stop' - Stop and unload the document\n\n"
                "🧭 Spatial commands: 'turn right', 'turn left', 'turn around'\n\n"
                "Keyboard: Spacebar to start, Escape to pause, F1 to stop.")
        return {"reply": intro}

    # 1️⃣3️⃣ DEFAULT: Check if it might be a filename
    # This handles cases where user just says the filename without "load"
    if not current_document_text:
        matched_file = find_matching_document(user_input)
        if matched_file:
            path = os.path.join(DOCUMENT_FOLDER, matched_file)
            full_text = read_document(path)
            
            if not full_text.startswith("Error"):
                current_document_text = full_text
                current_position = 0
                loaded_filename = matched_file
                current_document_name = os.path.splitext(matched_file)[0]
                current_document_type = os.path.splitext(matched_file)[1][1:].upper()
                
                return {"reply": f"Automatically loaded {current_document_name}. Say 'read' to begin, or 'summary' for summary."}

    # 1️⃣4️⃣ LLM FALLBACK for general questions
    try:
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": f"You are an assistant for the visually impaired. Be concise and helpful. Respond in {lang_name}."},
                {"role": "user", "content": request.text}
            ],
            temperature=0.1
        )

        return {"reply": completion.choices[0].message.content}

    except Exception as e:
        print("Error:", e)
        return {"reply": "Connection issue. Please try again."}

# -----------------------------
# HEALTH CHECK ENDPOINT
# -----------------------------
@app.get("/health")
async def health_check():
    files = [f for f in os.listdir(DOCUMENT_FOLDER) 
             if f.endswith(('.pdf', '.docx', '.txt'))] if os.path.exists(DOCUMENT_FOLDER) else []
    
    return {
        "status": "healthy",
        "documents_folder": DOCUMENT_FOLDER,
        "documents_count": len(files),
        "current_document": current_document_name if current_document_name else "None",
        "supported_formats": ["PDF", "DOCX", "TXT"]
    }

# -----------------------------
# RUN SERVER
# -----------------------------
if __name__ == "__main__":
    print("🚀 Starting LearnOutLoud Server...")
    print(f"📁 Using your Documents folder: {DOCUMENT_FOLDER}")
    print("📄 Supported formats: PDF, DOCX, TXT")
    print("🎤 Voice commands:")
    print("   • 'list documents' - Show all files")
    print("   • 'load filename' - Load a document")
    print("   • 'read' - Start reading")
    print("   • 'summary' - Get AI summary")
    print("   • 'find word' - Search in document")
    print("   • 'turn right/left' - Spatial navigation")
    print("🌐 Server running on http://localhost:8000")
    print("Press Ctrl+C to stop the server")
    uvicorn.run(app, host="0.0.0.0", port=8000)